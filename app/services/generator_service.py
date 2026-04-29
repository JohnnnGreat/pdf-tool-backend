"""Generator service — QR codes, barcodes, invoices, resumes, certificates."""
import io
import json
import os
from pathlib import Path


def generate_qr_code(data: str, output_path: str, size: int = 300, fill_color: str = "black", back_color: str = "white") -> None:
    import qrcode
    qr = qrcode.QRCode(box_size=10, border=4)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color=fill_color, back_color=back_color)
    img = img.resize((size, size))
    img.save(output_path)


def generate_barcode(data: str, output_path: str, barcode_type: str = "code128") -> None:
    import barcode as bc
    from barcode.writer import ImageWriter
    try:
        barcode_cls = bc.get_barcode_class(barcode_type)
    except Exception:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Unknown barcode type: {barcode_type}")
    bar = barcode_cls(data, writer=ImageWriter())
    stem = str(Path(output_path).with_suffix(""))
    bar.save(stem)


def generate_invoice(data: dict, output_path: str) -> None:
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.units import mm

    BLUE   = colors.HexColor("#2563EB")
    DARK   = colors.HexColor("#1e293b")
    MUTED  = colors.HexColor("#64748b")
    STRIPE = colors.HexColor("#f1f5f9")
    BORDER = colors.HexColor("#e2e8f0")

    currency     = data.get("currency", "$")
    tax_rate     = float(data.get("tax_rate", 0))
    discount_rate = float(data.get("discount", 0))

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )

    def ps(name, **kw):
        return ParagraphStyle(name, **kw)

    title_s     = ps("inv_title", fontName="Helvetica-Bold", fontSize=28, textColor=BLUE)
    label_s     = ps("inv_lbl",   fontName="Helvetica-Bold", fontSize=9,  textColor=MUTED)
    normal_s    = ps("inv_norm",  fontName="Helvetica",      fontSize=10, textColor=DARK, leading=15)
    small_s     = ps("inv_sm",    fontName="Helvetica",      fontSize=9,  textColor=MUTED, leading=13)
    right_s     = ps("inv_r",     fontName="Helvetica",      fontSize=10, textColor=DARK,  alignment=TA_RIGHT)
    right_bold_s= ps("inv_rb",    fontName="Helvetica-Bold", fontSize=11, textColor=DARK,  alignment=TA_RIGHT)
    th_s        = ps("inv_th",    fontName="Helvetica-Bold", fontSize=9,  textColor=colors.white)
    th_right_s  = ps("inv_thr",   fontName="Helvetica-Bold", fontSize=9,  textColor=colors.white, alignment=TA_RIGHT)

    story = []

    # ── Title + meta ──────────────────────────────────────────
    meta_inner = Table([
        [Paragraph("Invoice #", label_s), Paragraph(data.get("invoice_number", "INV-001"), normal_s)],
        [Paragraph("Date",      label_s), Paragraph(data.get("date", ""),           normal_s)],
        [Paragraph("Due Date",  label_s), Paragraph(data.get("due_date", "—"),      normal_s)],
    ], colWidths=[22*mm, 45*mm])
    meta_inner.setStyle(TableStyle([
        ("TOPPADDING",    (0,0),(-1,-1), 2),
        ("BOTTOMPADDING", (0,0),(-1,-1), 2),
    ]))

    header_row = Table([[Paragraph("INVOICE", title_s), meta_inner]], colWidths=["60%", "40%"])
    header_row.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP"), ("ALIGN",(1,0),(1,0),"RIGHT")]))
    story.append(header_row)
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE, spaceAfter=10, spaceBefore=8))

    # ── From / Bill To ────────────────────────────────────────
    from_text = (data.get("from", "") or "").replace("\n", "<br/>")
    to_text   = (data.get("to",   "") or "").replace("\n", "<br/>")
    addr_table = Table([
        [Paragraph("<b>FROM</b>",    label_s), Paragraph("<b>BILL TO</b>", label_s)],
        [Paragraph(from_text or "—", normal_s), Paragraph(to_text or "—", normal_s)],
    ], colWidths=["50%","50%"])
    addr_table.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1),2),
        ("BOTTOMPADDING",(0,0),(-1,-1),2),
    ]))
    story.append(addr_table)
    story.append(Spacer(1, 14))

    # ── Line items ────────────────────────────────────────────
    items = data.get("items", [])
    table_data = [[
        Paragraph("DESCRIPTION", th_s),
        Paragraph("QTY",         th_s),
        Paragraph("UNIT PRICE",  th_right_s),
        Paragraph("AMOUNT",      th_right_s),
    ]]
    subtotal = 0.0
    for i, item in enumerate(items):
        qty   = float(item.get("quantity", item.get("qty", 1)))
        price = float(item.get("unit_price", item.get("price", 0)))
        amt   = qty * price
        subtotal += amt
        bg = STRIPE if i % 2 else colors.white
        table_data.append([
            Paragraph(item.get("description", ""), normal_s),
            Paragraph(str(int(qty) if qty == int(qty) else qty), normal_s),
            Paragraph(f"{currency}{price:,.2f}", right_s),
            Paragraph(f"{currency}{amt:,.2f}",   right_s),
        ])

    items_table = Table(table_data, colWidths=["50%","10%","20%","20%"])
    row_bgs = [colors.white, STRIPE] * (len(items) // 2 + 1)
    items_table.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,0),  BLUE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.white, STRIPE]),
        ("GRID",          (0,0),(-1,-1), 0.25, BORDER),
        ("ALIGN",         (1,0),(-1,-1), "RIGHT"),
        ("TOPPADDING",    (0,0),(-1,-1), 7),
        ("BOTTOMPADDING", (0,0),(-1,-1), 7),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
        ("RIGHTPADDING",  (0,0),(-1,-1), 8),
    ]))
    story.append(items_table)
    story.append(Spacer(1, 6))

    # ── Totals ────────────────────────────────────────────────
    discount     = subtotal * discount_rate / 100
    after_disc   = subtotal - discount
    tax          = after_disc * tax_rate / 100
    grand_total  = after_disc + tax

    totals_rows = [[Paragraph("Subtotal", right_s), Paragraph(f"{currency}{subtotal:,.2f}", right_s)]]
    if discount_rate:
        totals_rows.append([Paragraph(f"Discount ({discount_rate:g}%)", right_s),
                            Paragraph(f"−{currency}{discount:,.2f}", right_s)])
    if tax_rate:
        totals_rows.append([Paragraph(f"Tax ({tax_rate:g}%)", right_s),
                            Paragraph(f"{currency}{tax:,.2f}", right_s)])
    totals_rows.append([Paragraph("<b>TOTAL</b>", right_bold_s),
                        Paragraph(f"<b>{currency}{grand_total:,.2f}</b>", right_bold_s)])

    totals_tbl = Table(totals_rows, colWidths=["70%","30%"])
    totals_tbl.setStyle(TableStyle([
        ("LINEABOVE",     (0,-1),(-1,-1), 1.5, BLUE),
        ("TOPPADDING",    (0,0),(-1,-1), 5),
        ("BOTTOMPADDING", (0,0),(-1,-1), 5),
    ]))
    story.append(totals_tbl)

    # ── Notes ─────────────────────────────────────────────────
    notes = (data.get("notes", "") or "").strip()
    if notes:
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=6))
        story.append(Paragraph("<b>Notes</b>", ps("inv_nh", fontName="Helvetica-Bold", fontSize=10, textColor=DARK)))
        story.append(Paragraph(notes.replace("\n", "<br/>"), small_s))

    doc.build(story)


def generate_invoice_xlsx(data: dict, output_path: str) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    BLUE   = "2563EB"
    LBLUE  = "EFF6FF"
    STRIPE = "F1F5F9"

    currency      = data.get("currency", "$")
    tax_rate      = float(data.get("tax_rate", 0))
    discount_rate = float(data.get("discount", 0))

    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 10
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16

    def cell(r, c, value="", bold=False, italic=False, size=10, color="1e293b",
             fill=None, align="left", num_fmt=None):
        cl = ws.cell(row=r, column=c, value=value)
        cl.font      = Font(name="Calibri", size=size, bold=bold, italic=italic, color=color)
        cl.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
        if fill:
            cl.fill = PatternFill(fill_type="solid", fgColor=fill)
        if num_fmt:
            cl.number_format = num_fmt
        return cl

    row = 1
    # Title
    cell(row, 1, "INVOICE", bold=True, size=24, color=BLUE)
    ws.row_dimensions[row].height = 36

    cell(row, 3, "Invoice #:", bold=True, size=10, align="right")
    cell(row, 4, data.get("invoice_number", "INV-001"), size=10)
    row += 1
    cell(row, 3, "Date:", bold=True, size=10, align="right")
    cell(row, 4, data.get("date", ""), size=10)
    row += 1
    cell(row, 3, "Due Date:", bold=True, size=10, align="right")
    cell(row, 4, data.get("due_date", ""), size=10)
    row += 2

    # From / To headers
    cell(row, 1, "FROM",    bold=True, size=9, color=BLUE)
    cell(row, 3, "BILL TO", bold=True, size=9, color=BLUE)
    row += 1

    from_lines = (data.get("from", "") or "").split("\n")
    to_lines   = (data.get("to",   "") or "").split("\n")
    for i in range(max(len(from_lines), len(to_lines), 1)):
        cell(row, 1, from_lines[i] if i < len(from_lines) else "")
        cell(row, 3, to_lines[i]   if i < len(to_lines)   else "")
        row += 1
    row += 1

    # Column headers
    for c, (hdr, algn) in enumerate([("Description","left"),("Qty","center"),("Unit Price","right"),("Amount","right")], 1):
        cell(row, c, hdr, bold=True, size=10, color="FFFFFF", fill=BLUE, align=algn)
    ws.row_dimensions[row].height = 20
    row += 1

    # Items
    items    = data.get("items", [])
    subtotal = 0.0
    for i, item in enumerate(items):
        qty   = float(item.get("quantity", item.get("qty",   1)))
        price = float(item.get("unit_price", item.get("price", 0)))
        amt   = qty * price
        subtotal += amt
        bg = STRIPE if i % 2 else "FFFFFF"
        cell(row, 1, item.get("description", ""), fill=bg)
        cell(row, 2, int(qty) if qty == int(qty) else qty, fill=bg, align="center")
        cell(row, 3, price, fill=bg, align="right", num_fmt=f'"{currency}"#,##0.00')
        cell(row, 4, amt,   fill=bg, align="right", num_fmt=f'"{currency}"#,##0.00')
        row += 1

    row += 1
    # Totals
    discount   = subtotal * discount_rate / 100
    after_disc = subtotal - discount
    tax        = after_disc * tax_rate / 100
    total      = after_disc + tax
    num = f'"{currency}"#,##0.00'

    cell(row, 3, "Subtotal:", align="right"); cell(row, 4, subtotal, align="right", num_fmt=num); row += 1
    if discount_rate:
        cell(row, 3, f"Discount ({discount_rate:g}%):", align="right")
        cell(row, 4, -discount, align="right", num_fmt=num); row += 1
    if tax_rate:
        cell(row, 3, f"Tax ({tax_rate:g}%):", align="right")
        cell(row, 4, tax, align="right", num_fmt=num); row += 1

    cell(row, 3, "TOTAL:", bold=True, size=12, align="right")
    cell(row, 4, total, bold=True, size=12, align="right", num_fmt=num)
    thin = Side(style="thin", color=BLUE)
    ws.cell(row, 3).border = Border(top=thin)
    ws.cell(row, 4).border = Border(top=thin)
    row += 2

    notes = (data.get("notes", "") or "").strip()
    if notes:
        cell(row, 1, "Notes", bold=True); row += 1
        cl = ws.cell(row=row, column=1, value=notes)
        cl.font      = Font(name="Calibri", size=10, color="64748b")
        cl.alignment = Alignment(wrap_text=True, vertical="top")
        ws.merge_cells(start_row=row, start_column=1, end_row=row+2, end_column=4)
        ws.row_dimensions[row].height = 48

    wb.save(output_path)


def generate_resume(data: dict, output_path: str) -> None:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib import colors
    from reportlab.lib.units import mm

    NAVY  = colors.HexColor("#1e3a5f")
    BLUE  = colors.HexColor("#2563EB")
    MUTED = colors.HexColor("#64748b")
    LIGHT = colors.HexColor("#e8f0fe")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm,
        topMargin=15*mm, bottomMargin=15*mm,
    )

    def ps(name, **kw):
        return ParagraphStyle(name, **kw)

    name_s    = ps("r_name", fontName="Helvetica-Bold", fontSize=22, textColor=NAVY, spaceAfter=2)
    contact_s = ps("r_ct",   fontName="Helvetica",      fontSize=9,  textColor=MUTED, leading=14)
    section_s = ps("r_sec",  fontName="Helvetica-Bold", fontSize=10, textColor=NAVY,
                   spaceBefore=10, spaceAfter=3, borderPadding=(0,0,2,0))
    normal_s  = ps("r_n",    fontName="Helvetica",      fontSize=9.5, textColor=colors.HexColor("#333333"), leading=14)
    bold_s    = ps("r_b",    fontName="Helvetica-Bold", fontSize=9.5, textColor=colors.HexColor("#111111"))
    date_s    = ps("r_d",    fontName="Helvetica",      fontSize=9,   textColor=MUTED, alignment=TA_RIGHT)
    skill_s   = ps("r_sk",   fontName="Helvetica",      fontSize=9.5, textColor=colors.HexColor("#333333"), leading=16)

    story = []

    # ── Header ────────────────────────────────────────────────
    contact_parts = [p for p in [
        data.get("email",""), data.get("phone",""),
        data.get("location",""), data.get("linkedin",""), data.get("website",""),
    ] if p]

    story.append(Paragraph(data.get("name", "Your Name"), name_s))
    if contact_parts:
        story.append(Paragraph(" · ".join(contact_parts), contact_s))
    story.append(HRFlowable(width="100%", thickness=2, color=NAVY, spaceAfter=8, spaceBefore=5))

    # ── Summary ───────────────────────────────────────────────
    summary = (data.get("summary") or "").strip()
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        story.append(Paragraph(summary, normal_s))

    # ── Experience ────────────────────────────────────────────
    work = data.get("work_experience", data.get("experience", []))
    if work:
        story.append(Paragraph("EXPERIENCE", section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        for job in work:
            title   = job.get("title", "")
            company = job.get("company", "")
            start   = job.get("start", "")
            end     = job.get("end", "Present")
            desc    = (job.get("description") or "").strip()

            row_tbl = Table([[
                Paragraph(f"<b>{title}</b>", bold_s),
                Paragraph(company, normal_s),
                Paragraph(f"{start} – {end}", date_s),
            ]], colWidths=["38%","38%","24%"])
            row_tbl.setStyle(TableStyle([
                ("VALIGN",(0,0),(-1,-1),"TOP"),
                ("TOPPADDING",(0,0),(-1,-1),0),
                ("BOTTOMPADDING",(0,0),(-1,-1),1),
            ]))
            story.append(row_tbl)
            for line in desc.split("\n"):
                line = line.strip().lstrip("•-").strip()
                if line:
                    story.append(Paragraph(f"• {line}", normal_s))
            story.append(Spacer(1, 5))

    # ── Education ─────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        for edu in education:
            institution = edu.get("institution", edu.get("school", ""))
            degree      = edu.get("degree", "")
            year        = edu.get("year", "")
            row_tbl = Table([[
                Paragraph(f"<b>{degree}</b>", bold_s),
                Paragraph(institution, normal_s),
                Paragraph(year, date_s),
            ]], colWidths=["40%","40%","20%"])
            row_tbl.setStyle(TableStyle([
                ("VALIGN",(0,0),(-1,-1),"TOP"),
                ("TOPPADDING",(0,0),(-1,-1),0),
                ("BOTTOMPADDING",(0,0),(-1,-1),1),
            ]))
            story.append(row_tbl)
            story.append(Spacer(1, 4))

    # ── Skills ────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        text = " · ".join(skills) if isinstance(skills, list) else str(skills)
        story.append(Paragraph(text, skill_s))

    # ── Certifications ────────────────────────────────────────
    for cert in data.get("certifications", []):
        if not story[-1].getPlainText() == "CERTIFICATIONS":  # type: ignore[attr-defined]
            story.append(Paragraph("CERTIFICATIONS", section_s))
            story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        if isinstance(cert, dict):
            story.append(Paragraph(
                f"• <b>{cert.get('name','')}</b> — {cert.get('issuer','')} {cert.get('year','')}".strip(" —"),
                normal_s))
        else:
            story.append(Paragraph(f"• {cert}", normal_s))

    # ── Languages ─────────────────────────────────────────────
    langs = data.get("languages", [])
    if langs:
        story.append(Paragraph("LANGUAGES", section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT, spaceAfter=4))
        story.append(Paragraph(" · ".join(langs) if isinstance(langs, list) else str(langs), skill_s))

    doc.build(story)


def generate_resume_docx(data: dict, output_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1e, 0x3a, 0x5f)
    MUTED = RGBColor(0x64, 0x74, 0x8b)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.8)
        sec.left_margin = sec.right_margin = Cm(2.2)

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def add_para(text="", bold=False, italic=False, size=10, color=None,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def add_section(title):
        p = add_para(title, bold=True, size=10, color=NAVY, space_before=10, space_after=2)
        # Bottom border via XML
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p.runs[0]._r.getparent().get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1e3a5f")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Name
    add_para(data.get("name", ""), bold=True, size=22, color=NAVY, space_after=2)

    contact_parts = [p for p in [
        data.get("email",""), data.get("phone",""),
        data.get("location",""), data.get("linkedin",""), data.get("website",""),
    ] if p]
    if contact_parts:
        add_para(" · ".join(contact_parts), size=9, color=MUTED, space_after=6)

    # Summary
    if (data.get("summary") or "").strip():
        add_section("PROFESSIONAL SUMMARY")
        add_para(data["summary"], size=10)

    # Experience
    work = data.get("work_experience", data.get("experience", []))
    if work:
        add_section("EXPERIENCE")
        for job in work:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(f"{job.get('title','')}  ·  {job.get('company','')}")
            r.bold = True; r.font.size = Pt(10)
            add_para(f"{job.get('start','')} – {job.get('end','Present')}", italic=True, size=9, color=MUTED, space_after=2)
            for line in (job.get("description") or "").split("\n"):
                line = line.strip().lstrip("•-").strip()
                if line:
                    bp = doc.add_paragraph(line, style="List Bullet")
                    bp.runs[0].font.size = Pt(9.5)

    # Education
    if data.get("education"):
        add_section("EDUCATION")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(f"{edu.get('degree','')}  ·  {edu.get('institution', edu.get('school',''))}")
            r.bold = True; r.font.size = Pt(10)
            if edu.get("year"):
                add_para(edu["year"], size=9, color=MUTED, space_after=2)

    # Skills
    skills = data.get("skills", [])
    if skills:
        add_section("SKILLS")
        add_para(", ".join(skills) if isinstance(skills, list) else str(skills), size=10)

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        add_section("CERTIFICATIONS")
        for cert in certs:
            if isinstance(cert, dict):
                add_para(f"{cert.get('name','')} — {cert.get('issuer','')} {cert.get('year','')}".strip(" —"), size=9.5)
            else:
                add_para(str(cert), size=9.5)

    # Languages
    langs = data.get("languages", [])
    if langs:
        add_section("LANGUAGES")
        add_para(", ".join(langs) if isinstance(langs, list) else str(langs), size=10)

    doc.save(output_path)


def generate_certificate(data: dict, output_path: str) -> None:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors

    doc = SimpleDocTemplate(output_path, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    center = ParagraphStyle("center", parent=styles["Normal"], alignment=TA_CENTER, fontSize=14)
    title_style = ParagraphStyle("title", parent=styles["Title"], alignment=TA_CENTER, fontSize=36)

    story = [
        Spacer(1, 60),
        Paragraph("Certificate of Achievement", title_style),
        Spacer(1, 30),
        Paragraph("This is to certify that", center),
        Spacer(1, 10),
        Paragraph(f"<b>{data.get('name', '')}</b>", ParagraphStyle("name", alignment=TA_CENTER, fontSize=28)),
        Spacer(1, 10),
        Paragraph(data.get("title", ""), center),
        Spacer(1, 20),
        Paragraph(f"Date: {data.get('date', '')}", center),
    ]
    doc.build(story)


def convert_color(value: str, source_format: str) -> dict:
    """Convert between HEX, RGB, and HSL."""
    if source_format == "hex":
        value = value.lstrip("#")
        r, g, b = int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)
    elif source_format == "rgb":
        parts = value.split(",")
        r, g, b = int(parts[0]), int(parts[1]), int(parts[2])
    else:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail="source_format must be 'hex' or 'rgb'")

    hex_val = f"#{r:02x}{g:02x}{b:02x}"
    r_, g_, b_ = r / 255, g / 255, b / 255
    cmax, cmin = max(r_, g_, b_), min(r_, g_, b_)
    delta = cmax - cmin
    l = (cmax + cmin) / 2
    s = 0 if delta == 0 else delta / (1 - abs(2 * l - 1))
    h = 0
    if delta:
        if cmax == r_:
            h = 60 * (((g_ - b_) / delta) % 6)
        elif cmax == g_:
            h = 60 * ((b_ - r_) / delta + 2)
        else:
            h = 60 * ((r_ - g_) / delta + 4)
    return {
        "hex": hex_val,
        "rgb": {"r": r, "g": g, "b": b},
        "hsl": {"h": round(h), "s": round(s * 100), "l": round(l * 100)},
    }


def generate_lorem_ipsum(count: int = 5, unit: str = "sentences") -> str:
    word_pool = (
        "lorem ipsum dolor sit amet consectetur adipiscing elit sed do eiusmod tempor "
        "incididunt ut labore et dolore magna aliqua enim ad minim veniam quis nostrud "
        "exercitation ullamco laboris nisi aliquip ex ea commodo consequat"
    ).split()
    import random
    words = word_pool[:]
    if unit == "words":
        return " ".join(random.choices(words, k=count))
    sentences = []
    for _ in range(count):
        n = random.randint(8, 16)
        s = " ".join(random.choices(words, k=n))
        sentences.append(s.capitalize() + ".")
    if unit == "sentences":
        return " ".join(sentences)
    paragraphs = []
    for _ in range(count):
        n = random.randint(3, 6)
        p = " ".join(
            " ".join(random.choices(words, k=random.randint(8, 16))).capitalize() + "."
            for _ in range(n)
        )
        paragraphs.append(p)
    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Favicon generator
# ---------------------------------------------------------------------------

_FAVICON_SIZES = [16, 32, 48, 64, 128, 180, 192, 512]
_ICO_SIZES = [(16, 16), (32, 32), (48, 48)]

# Font search paths (tried in order)
_FONT_PATHS = [
    "C:/Windows/Fonts/segoeui.ttf",
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/verdana.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    "/Library/Fonts/Arial.ttf",
]


def _get_font(size: int):
    from PIL import ImageFont
    for path in _FONT_PATHS:
        try:
            return ImageFont.truetype(path, size)
        except (OSError, IOError):
            continue
    return ImageFont.load_default()


def _hex_to_rgba(hex_color: str, alpha: int = 255) -> tuple:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    return (r, g, b, alpha)


def _draw_background(draw, size: int, bg_rgba: tuple, shape: str) -> None:
    from PIL import ImageDraw
    if shape == "circle":
        draw.ellipse([0, 0, size, size], fill=bg_rgba)
    elif shape == "rounded":
        radius = size // 5
        draw.rounded_rectangle([0, 0, size, size], radius=radius, fill=bg_rgba)
    else:
        draw.rectangle([0, 0, size, size], fill=bg_rgba)


def _save_favicon_pack(base_img, output_dir: str) -> list[str]:
    """Resize base_img to all favicon sizes and save PNG + ICO."""
    from PIL import Image
    paths = []

    # Individual PNGs
    for sz in _FAVICON_SIZES:
        resized = base_img.resize((sz, sz), Image.LANCZOS)
        path = os.path.join(output_dir, f"favicon-{sz}x{sz}.png")
        resized.save(path, "PNG")
        paths.append(path)

    # Multi-size ICO
    ico_path = os.path.join(output_dir, "favicon.ico")
    base_img.save(ico_path, format="ICO", sizes=_ICO_SIZES)
    paths.insert(0, ico_path)

    # web.manifest snippet
    manifest = {
        "icons": [
            {"src": f"favicon-{sz}x{sz}.png", "sizes": f"{sz}x{sz}", "type": "image/png"}
            for sz in _FAVICON_SIZES
        ]
    }
    manifest_path = os.path.join(output_dir, "site.webmanifest")
    with open(manifest_path, "w") as f:
        import json as _json
        _json.dump(manifest, f, indent=2)
    paths.append(manifest_path)

    return paths


def generate_favicon_from_text(
    text: str,
    bg_color: str = "#4F46E5",
    text_color: str = "#FFFFFF",
    shape: str = "square",
    output_dir: str = ".",
) -> list[str]:
    """Generate favicon pack from text or emoji.

    Args:
        text:       1–3 characters or an emoji shown on the icon
        bg_color:   background hex color, e.g. "#4F46E5"
        text_color: foreground hex color, e.g. "#FFFFFF"
        shape:      "square" | "circle" | "rounded"
        output_dir: directory to write files into

    Returns:
        List of file paths (favicon.ico, favicon-NxN.png …, site.webmanifest)
    """
    from PIL import Image, ImageDraw

    BASE = 512
    img = Image.new("RGBA", (BASE, BASE), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    bg_rgba = _hex_to_rgba(bg_color)
    fg_rgba = _hex_to_rgba(text_color)
    _draw_background(draw, BASE, bg_rgba, shape)

    # Auto-scale font to fit ~55 % of the canvas
    font_size = int(BASE * 0.55)
    font = _get_font(font_size)
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x = (BASE - tw) // 2 - bbox[0]
    y = (BASE - th) // 2 - bbox[1]
    draw.text((x, y), text, fill=fg_rgba, font=font)

    return _save_favicon_pack(img, output_dir)


def generate_favicon_from_image(input_path: str, output_dir: str) -> list[str]:
    """Generate favicon pack from an uploaded image.

    The image is center-cropped to a square before resizing.
    """
    from PIL import Image

    img = Image.open(input_path).convert("RGBA")
    w, h = img.size
    sq = min(w, h)
    left = (w - sq) // 2
    top = (h - sq) // 2
    img = img.crop((left, top, left + sq, top + sq))

    return _save_favicon_pack(img, output_dir)


def compute_file_hash(input_path: str) -> dict:
    import hashlib
    md5 = hashlib.md5()
    sha1 = hashlib.sha1()
    sha256 = hashlib.sha256()
    with open(input_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
            sha1.update(chunk)
            sha256.update(chunk)
    return {
        "md5": md5.hexdigest(),
        "sha1": sha1.hexdigest(),
        "sha256": sha256.hexdigest(),
        "file_size_bytes": os.path.getsize(input_path),
    }
